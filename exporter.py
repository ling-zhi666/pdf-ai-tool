# d:/Pdf_ai_tool/exporter.py
"""导出功能模块 - 支持导出到Excel和Word格式."""

import os
from datetime import datetime
from typing import List, Dict

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def can_export_excel():
    """检查是否可以导出Excel."""
    return HAS_OPENPYXL


def can_export_word():
    """检查是否可以导出Word."""
    return HAS_DOCX


def export_to_excel(records: List[Dict], file_path: str) -> bool:
    """
    导出记录到Excel文件.

    Args:
        records: 记录列表
        file_path: 输出文件路径

    Returns:
        bool: 是否导出成功
    """
    if not HAS_OPENPYXL:
        raise ImportError("openpyxl未安装，请运行: pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PDF摘要记录"

    # 表头样式
    header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="165DFF", end_color="165DFF", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")

    # 设置表头
    headers = ["ID", "文件名", "标签", "摘要", "导入时间"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    # 设置列宽
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 60
    ws.column_dimensions['E'].width = 18

    # 填充数据
    for row_idx, record in enumerate(records, 2):
        ws.cell(row=row_idx, column=1, value=record.get('id', ''))
        ws.cell(row=row_idx, column=2, value=record.get('file_name', ''))
        ws.cell(row=row_idx, column=3, value=record.get('tags', ''))

        # 摘要内容（处理None）
        summary = record.get('summary', '')
        if summary:
            # 去除emoji和格式符号，保留纯文本
            import re
            summary = re.sub(r'[\U0001F000-\U0001F9FF]', '', summary)  # 移除emoji
            summary = re.sub(r'[#📌📊📍💡📝]+', '', summary)
            summary = summary.strip()
        ws.cell(row=row_idx, column=4, value=summary or '')

        ws.cell(row=row_idx, column=5, value=record.get('create_time', ''))

    # 保存文件
    wb.save(file_path)
    return True


def export_to_word(records: List[Dict], file_path: str) -> bool:
    """
    导出记录到Word文件.

    Args:
        records: 记录列表
        file_path: 输出文件路径

    Returns:
        bool: 是否导出成功
    """
    if not HAS_DOCX:
        raise ImportError("python-docx未安装，请运行: pip install python-docx")

    doc = Document()

    # 设置文档标题
    title = doc.add_heading('PDF AI 摘要报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加生成时间
    time_para = doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    time_para.style = "Normal"

    # 添加统计
    doc.add_paragraph(f"共 {len(records)} 条记录")

    # 添加分隔线
    doc.add_paragraph("─" * 50)

    # 添加每条记录
    for record in records:
        # 文件名作为标题
        file_name = record.get('file_name', '未知文件')
        doc.add_heading(file_name, level=1)

        # 标签
        tags = record.get('tags', '')
        if tags:
            para = doc.add_paragraph(f"📌 标签: {tags}")
            para.style = "Normal"

        # 导入时间
        create_time = record.get('create_time', '')
        if create_time:
            para = doc.add_paragraph(f"📅 导入时间: {create_time}")
            para.style = "Normal"

        # 摘要内容
        summary = record.get('summary', '')
        if summary:
            summary_lines = summary.split('\n')
            for line in summary_lines:
                line = line.strip()
                if not line:
                    continue
                para = doc.add_paragraph(line)
                para.style = "Normal"

        # 分隔线
        doc.add_paragraph("─" * 30)

    # 保存文件
    doc.save(file_path)
    return True


def get_export_formats():
    """获取支持的导出格式."""
    formats = []
    if HAS_OPENPYXL:
        formats.append("Excel (.xlsx)")
    if HAS_DOCX:
        formats.append("Word (.docx)")
    return formats